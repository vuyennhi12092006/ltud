"""Tạo báo cáo Word theo mẫu Dự báo thời tiết.pdf"""
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "Bao_cao_Phan_mem_chup_man_hinh.docx")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.bold = bold
    if align:
        p.alignment = align
    return p


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Trang bìa
    for line, bold in [
        ("ĐẠI HỌC BÁCH KHOA HÀ NỘI", True),
        ("KHOA VẬT LÝ KỸ THUẬT", True),
        ("", False),
        ("----- □ & □ -----", False),
        ("", False),
        ("BÀI TẬP LỚN", True),
        ("MÔN LẬP TRÌNH ỨNG DỤNG", True),
        ("PH3460", True),
        ("", False),
        (
            "Đề tài: Xây dựng phần mềm chụp hình/quay màn hình máy tính.",
            False,
        ),
        ("", False),
        ("Sinh viên thực hiện: Nguyễn Thị Dung – 20227463", False),
        ("Nguyễn Thị Hương Ly – 20227473", False),
        ("", False),
        ("Giảng viên: Th.S Bùi Ngọc Hà", False),
        ("", False),
        ("Hà Nội, năm 2025", False),
    ]:
        p = add_para(doc, line, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # MỞ ĐẦU
    add_heading(doc, "MỞ ĐẦU", 1)
    intro = """Trong thời đại số, nhu cầu ghi lại, chia sẻ và xử lý nội dung hiển thị trên màn hình máy tính ngày càng phổ biến. Từ việc hướng dẫn sử dụng phần mềm, báo cáo lỗi hệ thống, lưu trữ tài liệu trực tuyến, cho đến tạo nội dung học tập và làm việc từ xa, chụp ảnh và quay video màn hình đã trở thành công cụ thiết yếu.

Các phần mềm chụp màn hình chuyên dụng như Snipping Tool, Lightshot, OBS hay ShareX cung cấp nhiều tính năng mạnh mẽ, tuy nhiên việc tự xây dựng một ứng dụng bằng Python giúp sinh viên nắm vững kiến thức lập trình ứng dụng, xử lý đồ họa, đa luồng, tích hợp AI (OCR) và kết nối dịch vụ đám mây.

Nhóm chúng em thực hiện đề tài "Xây dựng phần mềm chụp hình/quay màn hình máy tính" sử dụng ngôn ngữ Python với giao diện Tkinter. Phần mềm hỗ trợ chụp toàn màn hình hoặc vùng chọn, ghi chú ảnh, trích xuất văn bản bằng Tesseract OCR, quay video kèm âm thanh hệ thống và micro, quản lý thư viện ảnh/video, đồng thời upload lên Google Drive, AWS S3 hoặc Imgur.

Báo cáo gồm hai chương chính:
• Chương 1: Tổng quan
• Chương 2: Xây dựng và thử nghiệm chương trình

Kết quả dự kiến:
- Hoàn thiện phần mềm chụp và quay màn hình với giao diện thân thiện;
- Tích hợp OCR tự động copy văn bản vào clipboard;
- Hỗ trợ ghi chú ảnh và upload file lên đám mây;
- Ứng dụng chạy ổn định trên Windows.

Trong quá trình làm báo cáo, có thể còn sai sót. Chúng em mong thầy cô góp ý để hoàn thiện hơn. Xin chân thành cảm ơn!"""
    add_para(doc, intro)

    doc.add_page_break()
    add_heading(doc, "MỤC LỤC", 1)
    toc = """DANH MỤC HÌNH ẢNH
CHƯƠNG I. TỔNG QUAN
    1.1. Tổng quan về đề tài
    1.2. Công cụ thực hiện
        1.2.1. Phần cứng
        1.2.2. Phần mềm
    1.3. Xây dựng chương trình chụp/quay màn hình
        1.3.1. Lưu đồ thuật giải
        1.3.2. Phân tích thuật giải
CHƯƠNG II. XÂY DỰNG VÀ THỬ NGHIỆM CHƯƠNG TRÌNH
    2.1. Xây dựng hoàn chỉnh chương trình
    2.2. Thử nghiệm chương trình
KẾT LUẬN
TÀI LIỆU THAM KHẢO"""
    add_para(doc, toc)

    doc.add_page_break()
    add_heading(doc, "DANH MỤC HÌNH ẢNH", 1)
    add_para(doc, "Hình 1. Giao diện chính phần mềm chụp & ghi màn hình")
    add_para(doc, "Hình 2. Cửa sổ chọn vùng chụp màn hình")
    add_para(doc, "Hình 3. Công cụ ghi chú ảnh (Annotation Editor)")
    add_para(doc, "Hình 4. Kết quả OCR và thư viện ảnh/video")
    add_para(doc, "Hình 5. Ghi màn hình kèm âm thanh hệ thống và micro")

    doc.add_page_break()
    add_heading(doc, "CHƯƠNG I. TỔNG QUAN", 1)

    add_heading(doc, "1.1. Tổng quan về đề tài", 2)
    s11 = """Chụp và quay màn hình là nhu cầu thiết yếu trong học tập, làm việc và hỗ trợ kỹ thuật. Các công cụ thương mại cung cấp khả năng chụp nhanh, ghi chú, chia sẻ link, nhưng việc tự phát triển bằng Python giúp hiểu sâu cơ chế bắt pixel màn hình, xử lý ảnh, đa luồng và tích hợp API.

Phần mềm của nhóm được xây dựng với các module chính:
• main.py — giao diện Tkinter, chụp màn hình (mss), quay video (OpenCV), OCR, thư viện;
• audio_recorder.py — ghi âm thanh hệ thống (WASAPI loopback) và micro, ghép bằng ffmpeg;
• upload_service.py — upload Imgur, AWS S3, Google Drive và rút gọn link.

Các tính năng nổi bật: chụp toàn màn hình/vùng chọn, ghi chú (khung đỏ, mũi tên, chữ), OCR tiếng Việt/Anh, quay màn hình có tiếng, upload đám mây, quản lý gallery cục bộ."""
    add_para(doc, s11)

    add_heading(doc, "1.2. Công cụ thực hiện", 2)
    add_heading(doc, "1.2.1. Phần cứng", 3)
    add_para(
        doc,
        "Máy tính cá nhân (PC/Laptop) Windows 10 trở lên, RAM tối thiểu 8GB, CPU đa nhân; "
        "màn hình độ phân giải HD trở lên; micro và loa (cho ghi âm); kết nối Internet (upload/OCR).",
    )

    add_heading(doc, "1.2.2. Phần mềm", 3)
    s122 = """• Python 3.10+: ngôn ngữ lập trình chính;
• Tkinter: giao diện đồ họa (tích hợp sẵn Python);
• mss: chụp màn hình nhanh, đa màn hình;
• OpenCV (opencv-python): ghi video, xử lý khung hình;
• Pillow: xử lý và hiển thị ảnh;
• pytesseract + Tesseract OCR: trích xuất văn bản;
• sounddevice, pyaudiowpatch: ghi micro và âm thanh hệ thống;
• imageio-ffmpeg: ghép audio vào video;
• google-api-python-client, boto3, requests: upload đám mây.

Quy trình cài đặt:
1. Tạo môi trường ảo: python -m venv venv
2. Cài thư viện: pip install -r requirements.txt
3. Cài Tesseract OCR và cấu hình config.json
4. Chạy: python main.py"""
    add_para(doc, s122)

    add_heading(doc, "1.3. Xây dựng chương trình chụp/quay màn hình", 2)
    add_heading(doc, "1.3.1. Lưu đồ thuật giải", 3)
    add_para(doc, "Hình 1. Lưu đồ thuật giải (mô tả)")
    flow = """[Bắt đầu] → [Hiển thị giao diện chính + xem trước màn hình]
    → Người dùng chọn chức năng:
        • Chụp toàn màn hình / Chụp vùng chọn → Ẩn cửa sổ → Chụp → Annotation → Lưu gallery → OCR → Upload?
        • Ghi màn hình → Ghi video + audio song song → Dừng → Ghép ffmpeg → Upload?
        • Thư viện → Xem/Phát/OCR/Upload file đã lưu
    → [Kết thúc]"""
    add_para(doc, flow)

    add_heading(doc, "1.3.2. Phân tích thuật giải", 3)
    algo = """Khởi tạo: Import thư viện, đọc config.json, tạo thư mục gallery, khởi tạo giao diện Tkinter.

Chụp màn hình: Dùng mss.grab() lấy pixel màn hình. Chụp vùng: hiển thị RegionSelector toàn màn hình, người dùng kéo khung, crop và mở AnnotationEditor.

Ghi chú ảnh: Vẽ khung đỏ, mũi tên, chữ trên canvas, render vào PIL Image trước khi lưu.

OCR: Sau khi lưu ảnh, pytesseract.image_to_string(lang='vie+eng'), copy kết quả vào clipboard.

Quay video: Luồng phụ ghi khung hình bằng OpenCV VideoWriter; audio_recorder ghi WASAPI loopback + micro; khi dừng, ffmpeg ghép thành file MP4 cuối.

Upload: upload_service đọc upload_provider từ config, gọi API tương ứng, trả link (có thể rút gọn TinyURL).

Kết thúc: Giải phóng tài nguyên, đóng ứng dụng."""
    add_para(doc, algo)

    doc.add_page_break()
    add_heading(doc, "CHƯƠNG II. XÂY DỰNG VÀ THỬ NGHIỆM CHƯƠNG TRÌNH", 1)

    add_heading(doc, "2.1. Xây dựng hoàn chỉnh chương trình", 2)
    s21 = """(1) Khởi tạo và cấu hình

import tkinter as tk
import mss, cv2, pytesseract
from PIL import Image
from upload_service import load_config, upload_file

Đọc config.json cho Tesseract, upload provider, tùy chọn ghi âm.

(2) Giao diện chính

Các nút: Chụp toàn màn hình, Chụp vùng chọn, Ghi màn hình, Dừng ghi, Thư viện, Thoát.
Checkbox: Ghi âm thanh hệ thống, Ghi micro.
Khu vực xem trước màn hình real-time.

(3) Chụp và chỉnh sửa ảnh

Hàm grab_screen() và grab_region() dùng thư viện mss.
Sau chụp, AnnotationEditor cho phép vẽ khung/mũi tên/chữ.
Cửa sổ chính được ẩn (withdraw) trong suốt quá trình chụp và chỉnh sửa.

(4) OCR tự động

def run_ocr(image_path):
    text = pytesseract.image_to_string(Image.open(image_path), lang="vie+eng")
    pyperclip.copy(text)

(5) Ghi video kèm âm thanh

Video ghi vào file tạm *_video.mp4; audio ghi file *_audio.wav.
Khi dừng, mux_video_audio() ghép bằng ffmpeg thành file MP4 hoàn chỉnh.

(6) Upload đám mây

Trong upload_service.py, hàm upload_file() chọn provider:
- imgur: Client-ID
- s3: AWS access key, secret, bucket
- google_drive: service_account.json + folder_id"""
    add_para(doc, s21)

    add_heading(doc, "2.2. Thử nghiệm chương trình", 2)

    tests = [
        (
            "a) Thử nghiệm chụp toàn màn hình",
            "Chạy python main.py → bấm 'Chụp toàn màn hình' → chỉnh sửa → Hoàn tất.",
            "Ảnh lưu vào gallery/, OCR copy text, cửa sổ chính hiện lại sau khi hoàn tất.",
        ),
        (
            "b) Thử nghiệm chụp vùng chọn",
            "Bấm 'Chụp vùng chọn' → kéo khung → Hoàn tất.",
            "Không xuất hiện màn hình trắng cửa sổ chính khi đang chỉnh sửa; ảnh đúng vùng chọn.",
        ),
        (
            "c) Thử nghiệm OCR",
            "Chụp vùng có chữ tiếng Việt/Anh → kiểm tra clipboard.",
            "Văn bản được trích xuất và copy; hiện hộp thoại preview nội dung.",
        ),
        (
            "d) Thử nghiệm ghi màn hình + âm thanh",
            "Bật checkbox âm thanh HT và micro → Ghi màn hình → Dừng → phát video.",
            "Video MP4 có hình và tiếng; thời gian ghi hiển thị trên khung hình.",
        ),
        (
            "e) Thử nghiệm upload Google Drive",
            "Cấu hình google_drive trong config.json → Upload từ thư viện.",
            "Nhận link Google Drive, copy vào clipboard.",
        ),
    ]
    for title, steps, expected in tests:
        add_para(doc, title, bold=True)
        add_para(doc, f"• Thực hiện: {steps}")
        add_para(doc, f"• Kết quả mong đợi: {expected}")
        add_para(doc, "")

    add_para(
        doc,
        "Thử nghiệm toàn bộ đạt hiệu quả tốt: ứng dụng hoạt động trơn tru, "
        "các chức năng chính đáp ứng yêu cầu đề tài.",
    )

    doc.add_page_break()
    add_heading(doc, "KẾT LUẬN", 1)
    kl = """Báo cáo đã trình bày quá trình xây dựng phần mềm chụp hình/quay màn hình máy tính bằng Python. Ứng dụng sử dụng Tkinter, mss, OpenCV, Tesseract OCR và các dịch vụ upload đám mây, đáp ứng nhu cầu chụp, ghi chú, trích xuất văn bản, quay video có âm thanh và chia sẻ file.

Ưu điểm: giao diện đơn giản; tích hợp OCR và annotation; hỗ trợ nhiều nền tảng upload; mã nguồn module hóa dễ bảo trì.

Hạn chế: chất lượng OCR phụ thuộc Tesseract; ghi âm hệ thống cần Windows + pyaudiowpatch; upload Google Drive cần cấu hình Service Account.

Hướng phát triển: thêm phím tắt toàn cục, chọn monitor thứ hai, cải thiện OCR bằng AI, hỗ trợ upload trực tiếp qua OAuth Google cá nhân.

Qua bài tập lớn, nhóm đã củng cố kiến thức Python, lập trình GUI, xử lý ảnh, đa luồng và tích hợp API — nền tảng cho các dự án phần mềm ứng dụng thực tế."""
    add_para(doc, kl)

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "[1] Hà, B. V. (2023). Python cơ bản. ĐHQG Hà Nội.",
        "[2] Python Tkinter Documentation. https://docs.python.org/3/library/tkinter.html",
        "[3] MSS Documentation. https://python-mss.readthedocs.io/",
        "[4] Tesseract OCR. https://github.com/tesseract-ocr/tesseract",
        "[5] Google Drive API v3. https://developers.google.com/drive/api",
        "[6] OpenCV Python. https://docs.opencv.org/",
    ]
    for r in refs:
        add_para(doc, r)

    doc.save(OUTPUT)
    print("Created:", OUTPUT)


if __name__ == "__main__":
    main()
